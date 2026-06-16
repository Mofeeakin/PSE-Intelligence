"""
Build a professional DOCX compliance report matching the PSE company template standard.

Audit Report structure — exact replica of the Reflex Pay company template:
  Cover Page
  Document Information Table
  Table of Contents
  1.0 Executive Summary
  2.0 Audit Details  (2.1 Criteria | 2.2 Objectives | 2.3 Method | 2.4 Scope |
                      2.5 Findings Definition | 2.6 Opinion Rating Triangle)
  3.0 Audit Finding Definition — ALL core ISO clauses (4-10) merged into ONE
      section, each clause group rendered as a bold sub-divider + its own
      findings table (no page break between groups)
  5./6./7./8. ISMS (Annex A Control) Audit Findings — one section per Annex A
      group (Organisational/People/Physical/Technology), fixed-numbered
      regardless of which are present. NOTE: the template itself skips "4." —
      this numbering gap is intentional fidelity to the source document, not
      a bug.
  9.  Consolidated Audit Findings — comma-separated clause-ref lists per
      rating bucket (a. MaNC | b. MiNC | c. Observations | d. OFI), each
      bucket label is a real Heading-2 so it surfaces in the Word TOC field.
  10. Audit conclusions and audit recommendation — Yes/No checklist table.
  Colour is applied to the STATUS cell only; all other table cells are white.

Gap Assessment Report structure — exact replica of the Sterling company template:
  Cover Page
  Document Information Table
  Table of Contents
  Executive Summary
  Gap assessment Details (criteria | objectives | method | scope + info table)
  Gap assessment Findings Definition (Status) + CIRC legend (plain colour list)
  Gap Assessment Detailed Findings — Management Clauses: each core ISO clause
      (4-10) as its own Heading-3 + 6-column table (Clause No. | Clause Title |
      Detailed Control | Findings | Recommendations | Implementation Status)
  ISMS (Annex A Control) Findings — ONE continuous table for all 4 Annex A
      groups (5/6/7/8 in-table divider rows), not separate sections
  Summary of Findings — structured (Overall Outcome | Key Strengths | Main
      Gaps & Areas for Improvement | Conclusion) + ratings-distribution table
  Colour is applied to the Implementation Status cell only.

Ratings colour scheme:
  Audit  : A (green) | OFI/OBS (blue) | MiNC (amber) | MaNC (red)
  Gap    : Fully Implemented (green) | Partially Implemented (amber) |
           Not Implemented (red) | Not Applicable (grey)
"""
from __future__ import annotations

import io
import json
import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from reports.models import Report
from exports.colour_map import (
    AUDIT_RATING_COLOURS,
    GAP_RATING_COLOURS,
    get_colour,
    get_colour_by_status,
    infer_rating,
)

logger = logging.getLogger(__name__)

LOGO_PATH = Path(__file__).parent / "assets" / "pse_logo.png"
JSON_PREFIX = "__JSON__:"
SUMMARY_PREFIX = "__SUMMARY__:"  # gap-assessment "Summary of Findings" structured JSON tag

# Corporate colours
CORP_BLUE  = "002DA8"   # PSE corporate blue
DARK_NAVY  = "1E3A5F"   # headings / table headers
MID_GREY   = "666666"
LIGHT_BG   = "F4F6FA"   # alternating row tint (gap only)


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _hex_rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _style_run(run, size_pt: int = 10, bold: bool = False, colour: str = None, italic: bool = False):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if colour:
        run.font.color.rgb = _hex_rgb(colour)


_MD_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_MD_BOLD_ALT_RE = re.compile(r"__(.+?)__")
_MD_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)([^*\n]+?)\*(?!\*)")
_MD_HEADER_RE = re.compile(r"^#{1,6}\s+", re.MULTILINE)
_MD_BULLET_RE = re.compile(r"^[\-\*]\s+", re.MULTILINE)


def _strip_markdown(text: str) -> str:
    """Remove markdown formatting (**bold**, __bold__, *italic*, # headers, - bullets)
    that LLM output sometimes leaks despite prompt instructions — Word/PDF documents
    get their structure from real headings/styles, not text symbols."""
    text = _MD_BOLD_RE.sub(r"\1", text)
    text = _MD_BOLD_ALT_RE.sub(r"\1", text)
    text = _MD_ITALIC_RE.sub(r"\1", text)
    text = _MD_HEADER_RE.sub("", text)
    text = _MD_BULLET_RE.sub("", text)
    return text


def _sanitise(text) -> str:
    """Strip markdown formatting, replace § symbol with 'Clause ', collapse double spaces."""
    if not text:
        return ""
    if not isinstance(text, str):
        text = str(text)
    text = _strip_markdown(text)
    return text.replace("§", "Clause ").replace("  ", " ").strip()


# ── Section classifier (ISO 27001 core clause vs Annex A group) ──────────────
# Matches the literal `section` strings the wizard assigns (see
# Report_UI/src/lib/store.ts ISO_27001_QUESTIONS), so the audit report builder
# can replicate the Reflex Pay template's fixed structure: ISO core clauses
# 4-10 merge into one "3.0 Audit Finding Definition" section, while the four
# Annex A groups become fixed-numbered sections 5/6/7/8.

_ANNEX_GROUP_NUMBERS = {
    "organizational": 5,
    "organisational": 5,
    "people": 6,
    "physical": 7,
    "technology": 8,
}


def _classify_clause_section(section_name: str) -> tuple[str, "str | int"]:
    """
    Classify a ReportSection.section_name for the audit-report builder.
    Returns (kind, key):
      ("core",  section_name) — an ISO core clause (§4..§10) → merge into "3.0"
      ("annex", 5|6|7|8)      — an Annex A group → fixed-numbered section
      ("other", section_name) — unrecognised (other standards) → generic fallback
    """
    name = (section_name or "").strip()
    if name.startswith("§"):
        return "core", name
    if name.startswith("Annex A"):
        lower = name.lower()
        for keyword, num in _ANNEX_GROUP_NUMBERS.items():
            if keyword in lower:
                return "annex", num
    return "other", name


def _format_core_clause_label(section_name: str) -> str:
    """Turn '§4 Context' into 'Clause 4 — Context' for the bold sub-divider."""
    name = section_name.lstrip("§").strip()
    parts = name.split(" ", 1)
    if len(parts) == 2 and parts[0].replace(".", "").isdigit():
        return f"Clause {parts[0]} — {parts[1]}"
    return name


def _format_gap_clause_heading(section_name: str) -> str:
    """Turn '§4 Context' into 'Clause 4: Context' for a real Heading-3 (Sterling
    template style — colon, not em-dash, and an actual heading so it surfaces
    in the Word TOC at level 3)."""
    name = section_name.lstrip("§").strip()
    parts = name.split(" ", 1)
    if len(parts) == 2 and parts[0].replace(".", "").isdigit():
        return f"Clause {parts[0]}: {parts[1]}"
    return name


def _header_row(table, headers: list[str], bg: str = DARK_NAVY, font_size: int = 9):
    row = table.rows[0]
    for i, text in enumerate(headers):
        cell = row.cells[i]
        cell.text = text
        run = cell.paragraphs[0].runs[0]
        _style_run(run, size_pt=font_size, bold=True, colour="FFFFFF")
        _set_cell_bg(cell, bg)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)


def _cell_text(cell, text: str, size_pt: int = 9, bold: bool = False, colour: str = None):
    text = _sanitise(text)
    cell.text = text
    if cell.paragraphs[0].runs:
        run = cell.paragraphs[0].runs[0]
        _style_run(run, size_pt=size_pt, bold=bold, colour=colour)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)


def _add_heading(doc: Document, text: str, level: int = 1, colour: str = DARK_NAVY):
    h = doc.add_heading(_sanitise(text), level=level)
    if h.runs:
        h.runs[0].font.color.rgb = _hex_rgb(colour)
        h.runs[0].font.size = Pt(13 if level == 1 else 11)
    return h


# ── Cover Page ────────────────────────────────────────────────────────────────

def _add_cover(doc: Document, report: Report, is_gap: bool, logo=None):
    # Logo — use uploaded logo if provided, else fall back to static asset
    logo_path = None
    logo_width = 2.4
    if logo and logo.image:
        logo_path = logo.image.path
        logo_width = logo.width_inches
    elif LOGO_PATH.exists():
        logo_path = str(LOGO_PATH)

    if logo_path:
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Inches(logo_width))
            doc.add_paragraph()
        except Exception as exc:
            logger.warning("DOCX: Could not embed logo — %s", exc)

    # Report type badge
    badge_label = "GAP ASSESSMENT REPORT" if is_gap else "COMPLIANCE AUDIT REPORT"
    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = badge.add_run(badge_label)
    _style_run(r, size_pt=9, colour="9A9690")

    doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(report.title)
    _style_run(r, size_pt=24, bold=True, colour=CORP_BLUE)

    # Organisation
    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = org_p.add_run(report.organisation)
    _style_run(r, size_pt=14, colour="444444")

    doc.add_paragraph()

    meta = [
        "Standard: ISO/IEC 27001:2022",
        f"Report Type: {'Gap Assessment' if is_gap else 'Internal Audit'}",
    ]
    if report.department:
        meta.append(f"Department: {report.department}")

    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        _style_run(r, size_pt=11, colour=MID_GREY)

    doc.add_page_break()


def _apply_header_logo(doc: Document, logo) -> None:
    """Embed logo right-aligned in the page header of every section (every-page placement)."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        try:
            run.add_picture(logo.image.path, width=Inches(logo.width_inches))
        except Exception as exc:
            logger.warning("DOCX: Could not embed header logo — %s", exc)


def _add_section_break(doc: Document) -> None:
    """Insert a next-page section break so the new section gets its own header."""
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.oxml.ns import qn as _qn
    new_sect = _OxmlElement("w:sectPr")
    new_sect_type = _OxmlElement("w:type")
    new_sect_type.set(_qn("w:val"), "nextPage")
    new_sect.insert(0, new_sect_type)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(new_sect)


def _apply_header_logo_selected(doc: Document, logo) -> None:
    """
    Apply logo header only to sections whose 1-based index is in logo.pages.
    Sections are created by the caller via _add_section_break(); section index
    1 = cover, 2 = doc-info/TOC, 3+ = content.
    """
    selected = set(logo.pages or [])
    for i, section in enumerate(doc.sections, start=1):
        header = section.header
        header.is_linked_to_previous = False
        if i in selected:
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.clear()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run()
            try:
                run.add_picture(logo.image.path, width=Inches(logo.width_inches))
            except Exception as exc:
                logger.warning("DOCX: Could not embed selected-page header logo (section %d) — %s", i, exc)
        else:
            # Ensure this section has an empty header (don't inherit previous)
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.clear()


# ── Document Information Table ────────────────────────────────────────────────

def _add_doc_info_table(doc: Document, report: Report):
    _add_heading(doc, "Document Information", level=2)

    from datetime import date as _date
    author_name = ""
    if report.user:
        full = report.user.get_full_name()
        author_name = full if full.strip() else report.user.username

    std_code = report.standard.code if report.standard else "N/A"
    rows_data = [
        ("DOCUMENT REFERENCE", f"PSE-{report.id:04d}-{std_code}"),
        ("VERSION",            "1.0"),
        ("DATE OF CREATION",   _date.today().strftime("%d %B %Y")),
        ("DOCUMENT AUTHOR",    author_name or "PSE Consulting"),
        ("CLASSIFICATION",     "Internal — Confidential"),
        ("APPROVAL",           "Pending"),
    ]

    tbl = doc.add_table(rows=len(rows_data) + 1, cols=2)
    tbl.style = "Light Shading"
    _header_row(tbl, ["Field", "Value"], bg=CORP_BLUE)

    for i, (field, value) in enumerate(rows_data, start=1):
        _cell_text(tbl.rows[i].cells[0], field, bold=True)
        _cell_text(tbl.rows[i].cells[1], value)

    doc.add_paragraph()


# ── Table of Contents ─────────────────────────────────────────────────────────

def _insert_toc_field(doc: Document) -> None:
    """Insert a native Word TOC field (Heading 1–3) that Word rebuilds on F9."""
    p = doc.add_paragraph()
    p.clear()

    def _run_with_elem(elem_tag: str, **attrs):
        r = p.add_run()
        el = OxmlElement(elem_tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        r._r.append(el)
        return r

    _run_with_elem("w:fldChar", **{"w:fldCharType": "begin", "w:dirty": "true"})

    r_instr = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_instr._r.append(instr)

    _run_with_elem("w:fldChar", **{"w:fldCharType": "separate"})

    r_ph = p.add_run("(Open in Word and press F9 to generate page numbers)")
    _style_run(r_ph, size_pt=9, italic=True, colour=MID_GREY)

    _run_with_elem("w:fldChar", **{"w:fldCharType": "end"})


def _add_toc(doc: Document, report, is_gap: bool):
    """Insert a Table of Contents page using a native Word TOC field."""
    _add_heading(doc, "Table of Contents", level=1)
    _insert_toc_field(doc)
    doc.add_page_break()


# ── Rating Legend (gap assessment only) ──────────────────────────────────────

def _add_legend(doc: Document, service_type: str):
    """
    Gap assessment only — renders the CIRC rating legend matching Sterling's
    plain 1-column colour-filled list (no description column, no heading
    style — Sterling shows it as a bold lead-in paragraph, not a TOC heading).
    """
    label_p = doc.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = label_p.add_run("CIRC OPINION RATING")
    _style_run(r, size_pt=10, bold=True, colour=DARK_NAVY)

    # Only the three ratings actually assigned to findings (Fully/Partially/Not
    # Implemented) — Sterling's legend omits "Not Applicable".
    ordered_keys = ["fully_implemented", "partially_implemented", "not_implemented"]
    colours = GAP_RATING_COLOURS

    tbl = doc.add_table(rows=len(ordered_keys), cols=1)
    tbl.style = "Table Grid"

    for row_idx, key in enumerate(ordered_keys):
        meta = colours[key]
        cell = tbl.rows[row_idx].cells[0]
        cell.text = meta["label"]
        _set_cell_bg(cell, meta["bg"])
        if cell.paragraphs[0].runs:
            _style_run(cell.paragraphs[0].runs[0], size_pt=9, bold=True, colour=meta["fg"])

    doc.add_paragraph()


# ── Opinion Triangle (audit only — §2.6) ─────────────────────────────────────

def _draw_opinion_triangle() -> io.BytesIO:
    """
    Draw a colour-coded triangle pyramid in the style of the PSE audit template.
    Returns a PNG image in a BytesIO buffer.
    Layers (top to bottom):
        Green  — A / OFI      (conformant / opportunities)
        Amber  — OBS / MiNC   (observations / minor NC)
        Red    — MaNC          (major non-conformity)
    """
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        return None

    W, H = 520, 320
    img = Image.new("RGBA", (W, H), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    # Triangle points (apex at top-centre, base at bottom)
    apex_x = W // 2
    # Tier heights
    t1_y0, t1_y1 = 0,   H // 3         # top (green)
    t2_y0, t2_y1 = H // 3, 2 * H // 3  # mid (amber)
    t3_y0, t3_y1 = 2 * H // 3, H       # bottom (red)

    def _tier_poly(y0, y1, W, apex_x):
        """Trapezoid: narrower at top, wider at bottom."""
        total_H = H
        left0  = apex_x - (apex_x * y0 / total_H)
        right0 = apex_x + (apex_x * y0 / total_H)
        left1  = apex_x - (apex_x * y1 / total_H)
        right1 = apex_x + (apex_x * y1 / total_H)
        # Add padding so tier edges don't overlap
        pad = 2
        return [
            (left0 + pad, y0 + pad),
            (right0 - pad, y0 + pad),
            (right1 - pad, y1 - pad),
            (left1 + pad, y1 - pad),
        ]

    GREEN = (45, 106, 79)
    AMBER = (230, 126, 34)
    RED   = (192, 57, 43)

    draw.polygon(_tier_poly(t1_y0, t1_y1, W, apex_x), fill=GREEN)
    draw.polygon(_tier_poly(t2_y0, t2_y1, W, apex_x), fill=AMBER)
    draw.polygon(_tier_poly(t3_y0, t3_y1, W, apex_x), fill=RED)

    # Try to load a font; fall back to default
    font_sm, font_md, font_lg = None, None, None
    try:
        font_sm = ImageFont.truetype("arial.ttf", 16)
        font_md = ImageFont.truetype("arialbd.ttf", 22)
        font_lg = ImageFont.truetype("arialbd.ttf", 28)
    except Exception:
        try:
            font_sm = ImageFont.load_default()
            font_md = font_sm
            font_lg = font_sm
        except Exception:
            pass

    WHITE = (255, 255, 255)

    def _draw_centred(text, y_mid, font):
        if font is None:
            return
        bbox = draw.textbbox((0, 0), text, font=font)
        tw = bbox[2] - bbox[0]
        th = bbox[3] - bbox[1]
        draw.text(((W - tw) / 2, y_mid - th / 2), text, fill=WHITE, font=font)

    _draw_centred("A / OFI",      (t1_y0 + t1_y1) // 2, font_sm)
    _draw_centred("OBS / MiNC",   (t2_y0 + t2_y1) // 2, font_md)
    _draw_centred("MaNC",         (t3_y0 + t3_y1) // 2, font_lg)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ── Audit Details block (§2.0–§2.6) ─────────────────────────────────────────

def _add_audit_details(doc: Document, report: Report, scope_content: str):
    """
    Render the Audit Details section matching the PSE/Reflex Pay template structure.
    Includes: company table, criteria, objectives, method, scope, findings
    definition, and the opinion rating triangle.
    """
    _add_heading(doc, "2.0  Audit Details", level=1)

    from datetime import date as _date

    # Company info table ──────────────────────────────────────────────────────
    info_rows = [
        ("Company name:",      report.organisation or ""),
        ("Address:",           ""),
        ("Website:",           ""),
        ("",                   ""),
        ("Audit standard(s):", "ISO 27001:2022"),
        ("Date(s) of audit(s):", _date.today().strftime("%d %B %Y")),
        ("Duration:",          ""),
        ("Audit team members:", "PSE Consulting"),
        ("Scope Summary:",     (scope_content or report.scope or "")[:300]),
        ("Audit participants:", ""),
    ]

    tbl = doc.add_table(rows=len(info_rows), cols=2)
    tbl.style = "Table Grid"
    col_widths = [Cm(5.5), Cm(11.5)]
    for col_i, w in enumerate(col_widths):
        for cell in tbl.columns[col_i].cells:
            cell.width = w

    for row_i, (label, value) in enumerate(info_rows):
        lc = tbl.rows[row_i].cells[0]
        vc = tbl.rows[row_i].cells[1]
        lc.text = label
        vc.text = value
        if lc.paragraphs[0].runs:
            _style_run(lc.paragraphs[0].runs[0], size_pt=9, bold=True)
        if vc.paragraphs[0].runs:
            _style_run(vc.paragraphs[0].runs[0], size_pt=9)

    doc.add_paragraph()

    # §2.1 Audit Criteria ─────────────────────────────────────────────────────
    _add_heading(doc, "2.1  Audit Criteria and Reference Documents", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "The international standard for information security ISO/IEC 27001:2022 (and its "
        "subsequent revisions) will be used as the basis of the criteria for the audit programme. "
        "The audit was carried out in English language."
    )
    _style_run(r, size_pt=10)
    doc.add_paragraph()

    # §2.2 Audit Objectives ───────────────────────────────────────────────────
    _add_heading(doc, "2.2  Audit Objectives", level=2)
    objectives_intro = doc.add_paragraph()
    objectives_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = objectives_intro.add_run(
        "In line with the requirements of the international standards, "
        "the overall objectives of this internal audit are to:"
    )
    _style_run(r, size_pt=10)

    objectives = [
        "Ensure that the implemented management systems and best practices are carried out "
        "effectively, efficiently, and economically to benefit the organisation.",
        "Determine the conformity of the organisation's management systems with audit criteria.",
        "Identify further opportunities for continual improvement, which may extend beyond the "
        "criteria set out in international standards.",
        "Provide the organisation with the internal assurance that the management systems are "
        "effectively managed and risks to the business are minimised.",
    ]
    for obj in objectives:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        r = p.add_run(obj)
        _style_run(r, size_pt=10)
    doc.add_paragraph()

    # §2.3 Audit Method ───────────────────────────────────────────────────────
    _add_heading(doc, "2.3  Audit Method", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "The audit was conducted in accordance with ISO 19011. The combination of checklists "
        "and questionnaires, interviews, observation, documents, system, and record review with "
        "the auditee was used to collect evidence to validate the effectiveness of the implemented "
        "Information Security Management System (ISMS)."
    )
    _style_run(r, size_pt=10)
    doc.add_paragraph()

    # §2.4 Scope of Audit ─────────────────────────────────────────────────────
    _add_heading(doc, "2.4  Scope of Audit", level=2)
    scope_text = scope_content or report.scope or \
        "The audit covered the organisation's Information Security Management System (ISMS) " \
        "against the requirements of ISO/IEC 27001:2022, including all in-scope locations, " \
        "business activities, processes, information assets, and products and services."
    for line in scope_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line)
        _style_run(r, size_pt=10)
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()

    # §2.5 Audit Findings Definition ──────────────────────────────────────────
    _add_heading(doc, "2.5  Audit Findings Definition (Status)", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "Where a discrepancy against the standard has been found, one of five types of items "
        "has been raised as follows:"
    )
    _style_run(r, size_pt=10)

    findings_defs = [
        ("Acceptable (A):", "conformance with the management system or process."),
        ("Major Non-Conformity (MaNC):", "a significant issue that represents a breakdown of the "
         "operation of the management system or process."),
        ("Minor Non-Conformity (MiNC):", "a single lapse that does not indicate a breakdown of the "
         "management system or process."),
        ("Observation (OBS):", "a comment which may be of use to the auditee based on the experience "
         "of other implementations and the expertise of the internal auditor."),
        ("Opportunity for Improvement (OFI):", "an observation or suggestion regarding a potential "
         "improvement opportunity. No action is necessarily required."),
    ]
    for label, definition in findings_defs:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        r_bold = p.add_run(label + " ")
        _style_run(r_bold, size_pt=10, bold=True)
        r_def = p.add_run(definition)
        _style_run(r_def, size_pt=10)
    doc.add_paragraph()

    # §2.6 Opinion Rating ─────────────────────────────────────────────────────
    _add_heading(doc, "2.6  Opinion Rating (1-5)", level=2)
    triangle_buf = _draw_opinion_triangle()
    if triangle_buf:
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        run.add_picture(triangle_buf, width=Inches(3.5))
    else:
        # Fallback: small text table
        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = "Light Shading"
        _header_row(tbl, ["Rating", "Description"])
        fallback = [
            ("A / OFI",    "Acceptable / Opportunity for Improvement"),
            ("OBS / MiNC", "Observation / Minor Non-Conformity"),
            ("MaNC",       "Major Non-Conformity"),
        ]
        colours_fb = ["1D6A3C", "E67E22", "C0392B"]
        for i, (lbl, dsc) in enumerate(fallback, start=1):
            _cell_text(tbl.rows[i].cells[0], lbl, bold=True, colour="FFFFFF")
            _set_cell_bg(tbl.rows[i].cells[0], colours_fb[i - 1])
            _cell_text(tbl.rows[i].cells[1], dsc)
    doc.add_paragraph()
    doc.add_page_break()


# ── Gap Assessment Details (Sterling template equivalent of "2.0 Audit Details") ─

def _add_gap_assessment_details(doc: Document, report: Report, scope_content: str):
    """
    Render the "Gap assessment Details" section matching the Sterling template
    structure. Like _add_audit_details(), this is mostly static boilerplate
    text (criteria/objectives/method are fixed wording) — only the info table
    and the Scope sub-section are dynamic, reusing the already-generated
    scope_content the same way _add_audit_details() does.
    """
    _add_heading(doc, "Gap assessment Details", level=1)

    # Info table — Sterling's exact 5-row layout ───────────────────────────────
    info_rows = [
        ("DATE(S)",                     ""),
        ("LOCATION(S)",                 ""),
        ("GAP ASSESSMENTOR(S)",         "PSE Consulting"),
        ("GAP ASSESSMENT PARTICIPANTS", ""),
        ("SCOPE SUMMARY",               (scope_content or report.scope or "")[:300]),
    ]

    tbl = doc.add_table(rows=len(info_rows), cols=2)
    tbl.style = "Table Grid"
    col_widths = [Cm(5.5), Cm(11.5)]
    for col_i, w in enumerate(col_widths):
        for cell in tbl.columns[col_i].cells:
            cell.width = w

    for row_i, (label, value) in enumerate(info_rows):
        lc = tbl.rows[row_i].cells[0]
        vc = tbl.rows[row_i].cells[1]
        lc.text = label
        vc.text = value
        if lc.paragraphs[0].runs:
            _style_run(lc.paragraphs[0].runs[0], size_pt=9, bold=True)
        if vc.paragraphs[0].runs:
            _style_run(vc.paragraphs[0].runs[0], size_pt=9)

    doc.add_paragraph()

    # Gap Assessment Criteria and Reference Documents ──────────────────────────
    _add_heading(doc, "Gap Assessment Criteria and Reference Documents", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "The international standard for ISO/IEC 27001 (Information security) (and its "
        "subsequent revisions) was used as the basis of the criteria for the gap assessment "
        "programme. The gap assessment was carried out in English language."
    )
    _style_run(r, size_pt=10)
    doc.add_paragraph()

    # Gap Assessment Objectives ──────────────────────────────────────────────────
    _add_heading(doc, "Gap Assessment Objectives", level=2)
    objectives_intro = doc.add_paragraph()
    objectives_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = objectives_intro.add_run(
        "In line with the requirements of the international standards, "
        "the overall objectives of this internal gap assessment are to:"
    )
    _style_run(r, size_pt=10)

    objectives = [
        "Ensure that the implemented management systems and best practices are carried out "
        "effectively, efficiently, and economically to benefit the organisation.",
        "Determine the conformity of the organisation's management systems with gap assessment "
        "criteria.",
        "Identify further opportunities for continual improvement, which may extend beyond the "
        "criteria set out in international standards.",
        "Provide the organisation with the internal assurance that the management systems are "
        "effectively managed and risks to the business are minimised.",
    ]
    for obj in objectives:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        r = p.add_run(obj)
        _style_run(r, size_pt=10)
    doc.add_paragraph()

    # Gap Assessment Method ──────────────────────────────────────────────────────
    _add_heading(doc, "Gap Assessment Method", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "The gap assessment was conducted physically and virtually. The combination of checklists "
        "and questionnaires, interviews, observation, documents, system, and record review with "
        "the auditee was used to collect evidence to validate the effectiveness of the implemented "
        "Information Security Management System (ISMS)."
    )
    _style_run(r, size_pt=10)
    doc.add_paragraph()

    # Scope of Gap Assessment ─────────────────────────────────────────────────────
    _add_heading(doc, "Scope of Gap Assessment", level=2)
    scope_text = scope_content or report.scope or \
        "The gap assessment covered the organisation's Information Security Management System " \
        "(ISMS) against the requirements of ISO/IEC 27001:2022, including all in-scope locations, " \
        "business activities, processes, information assets, and products and services."
    for line in scope_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line)
        _style_run(r, size_pt=10)
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()
    doc.add_page_break()


def _add_gap_findings_definition(doc: Document):
    """
    Render "Gap assessment Findings Definition (Status)" matching Sterling's
    3-line CIRC definition, followed by the simplified colour-filled legend.
    """
    _add_heading(doc, "Gap assessment Findings Definition (Status)", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(
        "Where a discrepancy against the standard has been found, one of three types of items "
        "has been raised as follows:"
    )
    _style_run(r, size_pt=10)

    findings_defs = [
        ("Fully Implemented:", "conformance with the management system or process."),
        ("Not Implemented:", "a significant issue that represents a breakdown of the operation "
         "of the management system or process."),
        ("Partially Implemented:", "a single lapse that does not indicate a breakdown of the "
         "management system or process."),
    ]
    for label, definition in findings_defs:
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        r_bold = p.add_run(label + " ")
        _style_run(r_bold, size_pt=10, bold=True)
        r_def = p.add_run(definition)
        _style_run(r_def, size_pt=10)
    doc.add_paragraph()

    _add_legend(doc, "gap_assessment")
    doc.add_page_break()


# ── Consolidated Audit Findings (§9) ─────────────────────────────────────────

def _add_consolidated_findings(doc: Document, all_findings: list[dict]):
    """
    Render a consolidated grouped summary of all clause findings — matching
    the Reflex Pay template's "9. Consolidated Audit Findings" layout exactly:
    each rating bucket is a real Heading-2 ("a. Major Non-conformities [MaNC]"
    etc, so it surfaces in the Word TOC field), followed by a single plain
    paragraph listing the affected clause references as a comma-separated list
    (not full finding sentences) — e.g. "Clause 5.2, Annex 5.30, Annex A.8.6".
    """
    _add_heading(doc, "9.  Consolidated Audit Findings", level=1)

    buckets = {
        "MaNC": [],
        "MiNC": [],
        "OBS":  [],
        "OFI":  [],
    }
    # Accept both full names and abbreviations
    alias = {
        "Major Non-Conformity": "MaNC",
        "Major Non-Conformance": "MaNC",
        "Minor Non-Conformity": "MiNC",
        "Minor Non-Conformance": "MiNC",
        "Observation": "OBS",
        "Opportunity for Improvement": "OFI",
    }
    for f in all_findings:
        raw = (f.get("status") or "").strip()
        key = alias.get(raw, raw)
        if key in buckets:
            buckets[key].append(f)

    bucket_order = [
        ("a", "MaNC", "Major Non-conformities [MaNC]"),
        ("b", "MiNC", "Minor Non-conformities [MiNC]"),
        ("c", "OBS",  "Observations"),
        ("d", "OFI",  "Opportunities for Improvement"),
    ]

    for letter, key, label in bucket_order:
        _add_heading(doc, f"{letter}. {label}", level=2, colour=DARK_NAVY)

        clause_refs = [f.get("clause_ref", "").strip() for f in buckets[key] if f.get("clause_ref")]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if clause_refs:
            r = p.add_run(", ".join(f"Clause {c}" for c in clause_refs))
            _style_run(r, size_pt=10)
        else:
            r = p.add_run("None identified.")
            _style_run(r, size_pt=10, colour=MID_GREY)
        doc.add_paragraph()


# ── Audit Conclusions checklist (§10) ────────────────────────────────────────

def _add_audit_conclusions(doc: Document):
    """
    Render the §10 Audit Conclusions table matching the PSE template:
    3-row Yes/No checklist table.
    """
    _add_heading(doc, "10.  Audit conclusions and audit recommendation", level=1)

    # Introductory checklist table
    checklist = [
        ("Has there been any serious deviation from the audit plan? (If yes, please specify)", "No ✓"),
        ("The management review processes are in place and adequate", "Yes ✓"),
        ("The audit was successful in meeting the stated objectives", "Yes ✓"),
    ]

    tbl = doc.add_table(rows=len(checklist), cols=2)
    tbl.style = "Table Grid"
    col_widths = [Cm(13.5), Cm(3.5)]
    for col_i, w in enumerate(col_widths):
        for cell in tbl.columns[col_i].cells:
            cell.width = w

    for row_i, (question, answer) in enumerate(checklist):
        qc = tbl.rows[row_i].cells[0]
        ac = tbl.rows[row_i].cells[1]
        qc.text = question
        ac.text = answer
        if qc.paragraphs[0].runs:
            _style_run(qc.paragraphs[0].runs[0], size_pt=10)
        if ac.paragraphs[0].runs:
            _style_run(ac.paragraphs[0].runs[0], size_pt=10, bold=True)
        qc.paragraphs[0].paragraph_format.space_before = Pt(3)
        qc.paragraphs[0].paragraph_format.space_after = Pt(3)

    doc.add_paragraph()


# ── Summary of Findings (gap assessment only — Sterling template) ────────────

def _add_gap_summary_of_findings(doc: Document, section, all_clause_findings: list[dict]):
    """
    Render the "Summary of Findings" section matching the Sterling template:
    Overall Outcome / Key Strengths / Main Gaps & Areas for Improvement /
    Conclusion as bold-run paragraph labels each followed by a bullet list
    (Conclusion gets prose instead), plus a ratings-distribution count table.

    `section` is the ReportSection whose content is SUMMARY_PREFIX-tagged JSON
    (see agents/pipeline.py._generate_conclusion). Falls back to plain prose
    rendering if the JSON is missing/malformed, so a bad LLM response never
    crashes the export.
    """
    _add_heading(doc, "Summary of Findings", level=1)

    content = (section.content if section else "") or ""
    data = None
    if content.startswith(SUMMARY_PREFIX):
        raw_json = content[len(SUMMARY_PREFIX):]
        try:
            data = json.loads(raw_json)
        except json.JSONDecodeError:
            logger.warning("DOCX: Failed to parse Summary of Findings JSON — falling back to prose.")

    def _bullet_block(label: str, items: list):
        p_label = doc.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p_label.add_run(label)
        _style_run(r, size_pt=10, bold=True)
        for item in items or []:
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            r = p.add_run(_sanitise(item))
            _style_run(r, size_pt=10)
        doc.add_paragraph()

    if isinstance(data, dict):
        _bullet_block("Overall Outcome", data.get("overall_outcome"))
        _bullet_block("Key Strengths", data.get("key_strengths"))
        _bullet_block("Main Gaps & Areas for Improvement", data.get("main_gaps"))

        p_label = doc.add_paragraph()
        r = p_label.add_run("Conclusion")
        _style_run(r, size_pt=10, bold=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(_sanitise(data.get("conclusion") or ""))
        _style_run(r, size_pt=10)
        doc.add_paragraph()
    else:
        # Fallback: render whatever text we have as plain prose.
        _add_prose_section_body(doc, content)

    # Ratings-distribution count table — computed directly from findings, no LLM.
    counts = {"Fully Implemented": 0, "Partially Implemented": 0, "Not Implemented": 0}
    label_map = {
        "FI": "Fully Implemented", "Fully Implemented": "Fully Implemented",
        "PI": "Partially Implemented", "Partially Implemented": "Partially Implemented",
        "NI": "Not Implemented", "Not Implemented": "Not Implemented",
    }
    for f in all_clause_findings:
        key = label_map.get((f.get("circ_rating") or "").strip())
        if key:
            counts[key] += 1

    tbl = doc.add_table(rows=len(counts), cols=2)
    tbl.style = "Table Grid"
    for row_i, (label, count) in enumerate(counts.items()):
        colour = get_colour_by_status(label, "gap_assessment")
        lc, vc = tbl.rows[row_i].cells[0], tbl.rows[row_i].cells[1]
        _cell_text(lc, label, size_pt=9, bold=True, colour=colour["fg"])
        _set_cell_bg(lc, colour["bg"])
        _cell_text(vc, str(count), size_pt=9, bold=True)
        _set_cell_bg(vc, "FFFFFF")

    doc.add_paragraph()


# ── Score Summary ─────────────────────────────────────────────────────────────
# NOTE: not currently called by build_docx() — the Sterling template has no
# separate "Compliance Score Summary" section for gap assessments (the
# ratings-distribution table inside _add_gap_summary_of_findings() covers
# the equivalent role). Kept for potential reuse.

def _add_score_summary(doc: Document, report: Report):
    score = getattr(report, "compliance_score", None)
    if not score:
        return

    _add_heading(doc, "Compliance Score Summary", level=2)

    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Light Shading"
    _header_row(tbl, ["Dimension", "Score", "Weight"])

    def _fmt(v) -> str:
        return f"{float(v):.1f}%" if v is not None else "N/A"

    rows_data = [
        ("Section Score",     _fmt(score.section_score),     "40%"),
        ("Evidence Score",    _fmt(score.evidence_score),    "35%"),
        ("Consistency Score", _fmt(score.consistency_score), "25%"),
    ]
    for i, (dim, val, wt) in enumerate(rows_data, start=1):
        _cell_text(tbl.rows[i].cells[0], dim)
        _cell_text(tbl.rows[i].cells[1], val, bold=True)
        _cell_text(tbl.rows[i].cells[2], wt)

    doc.add_paragraph()
    total_p = doc.add_paragraph()
    total_str = _fmt(score.total_score)
    r = total_p.add_run(f"Overall Score: {total_str}  |  Status: {score.status or 'N/A'}")
    _style_run(r, size_pt=13, bold=True, colour=DARK_NAVY)
    doc.add_paragraph()


# ── Prose section renderer ────────────────────────────────────────────────────

def _add_prose_section_body(doc: Document, content: str):
    """Render prose paragraphs without adding a heading (caller adds its own)."""
    for line in content.split("\n"):
        line = _sanitise(line.strip())
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(line)
        _style_run(r, size_pt=10)
        p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()


def _add_prose_section(doc: Document, section_name: str, content: str):
    _add_heading(doc, _sanitise(section_name), level=1)
    _add_prose_section_body(doc, content)


# ── Clause findings table (JSON section renderer) ────────────────────────────

def _add_audit_findings_table_body(doc: Document, findings: list[dict]):
    """
    Render the 5-column colour-coded findings table without adding a heading
    (caller adds its own heading or bold divider).
    Columns: CONTROL | AUDIT FINDINGS | STATUS | EVIDENCE REVIEWED | RECOMMENDATION
    Only the STATUS cell is colour-coded; all other cells have white background.
    """
    headers = ["CONTROL", "AUDIT FINDINGS", "STATUS", "EVIDENCE REVIEWED", "RECOMMENDATION"]
    col_widths = [Cm(3.5), Cm(5.0), Cm(2.5), Cm(4.5), Cm(4.5)]

    tbl = doc.add_table(rows=1 + len(findings), cols=len(headers))
    tbl.style = "Table Grid"
    _header_row(tbl, headers, bg=DARK_NAVY)

    # Set column widths
    for col_i, w in enumerate(col_widths):
        for cell in tbl.columns[col_i].cells:
            cell.width = w

    for row_idx, f in enumerate(findings, start=1):
        cells = tbl.rows[row_idx].cells
        status_str = f.get("status", "OBS")
        colour = get_colour_by_status(status_str, "audit_report")

        clause         = f.get("clause_ref") or ""
        control        = (f.get("requirement_summary") or f.get("control_name") or "")[:100]
        finding        = (f.get("audit_finding") or "")[:250]
        evidence       = (f.get("evidence_reviewed") or f.get("evidence") or "")[:200]
        recommendation = (f.get("recommendation") or "")[:200]

        _cell_text(cells[0], f"{clause} {control}".strip(), size_pt=9, bold=True)
        _cell_text(cells[1], finding,                       size_pt=9)
        _cell_text(cells[2], colour["label"],               size_pt=9, bold=True, colour=colour["fg"])
        _set_cell_bg(cells[2], colour["bg"])
        _cell_text(cells[3], evidence,                      size_pt=9)
        _cell_text(cells[4], recommendation,                size_pt=9)

        # Explicit white background for non-status cells
        for ci in (0, 1, 3, 4):
            _set_cell_bg(cells[ci], "FFFFFF")

    doc.add_paragraph()


def _add_audit_findings_table(doc: Document, section_name: str, findings: list[dict]):
    """
    Render audit findings as a colour-coded 5-column table matching the PSE template,
    with its own Heading-2 section title. Used for the generic ("other"-standard) path.
    """
    _add_heading(doc, section_name, level=2, colour=CORP_BLUE)
    _add_audit_findings_table_body(doc, findings)


def _add_gap_findings_table(doc: Document, section_name: str, findings: list[dict]):
    """Render gap findings as a colour-coded 6-column table."""
    _add_heading(doc, section_name, level=2, colour=CORP_BLUE)

    headers = ["Clause Ref", "Control", "CIRC Rating", "Current State", "Required State", "Recommendation"]
    tbl = doc.add_table(rows=1 + len(findings), cols=len(headers))
    tbl.style = "Table Grid"
    _header_row(tbl, headers, bg=DARK_NAVY)

    for row_idx, f in enumerate(findings, start=1):
        cells = tbl.rows[row_idx].cells
        rating_str = f.get("circ_rating", "Not Implemented")
        colour = get_colour_by_status(rating_str, "gap_assessment")

        _cell_text(cells[0], f.get("clause_ref") or "",              size_pt=9, bold=True)
        _cell_text(cells[1], (f.get("control_name") or "")[:80],    size_pt=9)
        _cell_text(cells[2], colour["label"],                        size_pt=9, bold=True, colour=colour["fg"])
        _set_cell_bg(cells[2], colour["bg"])
        _cell_text(cells[3], (f.get("current_state") or "")[:200],  size_pt=9)
        _cell_text(cells[4], (f.get("required_state") or "")[:150], size_pt=9)
        _cell_text(cells[5], f.get("recommendation") or "",          size_pt=9)

        # Explicit white background for non-rating cells (avoid Word table-style banding)
        for ci in (0, 1, 3, 4, 5):
            _set_cell_bg(cells[ci], "FFFFFF")

    doc.add_paragraph()


def _add_gap_clause_table_body(doc: Document, findings: list[dict]):
    """
    Render a per-clause gap findings table matching Sterling's exact 6-column
    layout: Clause No. | Clause Title | Detailed Control | Findings |
    Recommendations | Implementation Status. No heading (caller adds the
    Heading-3 clause divider). Only the Implementation Status cell is
    colour-coded; all other cells are white.
    """
    headers = ["Clause No.", "Clause Title", "Detailed Control", "Findings",
               "Recommendations", "Implementation Status"]
    tbl = doc.add_table(rows=1 + len(findings), cols=len(headers))
    tbl.style = "Table Grid"
    _header_row(tbl, headers, bg=DARK_NAVY)

    for row_idx, f in enumerate(findings, start=1):
        cells = tbl.rows[row_idx].cells
        rating_str = f.get("circ_rating", "Not Implemented")
        colour = get_colour_by_status(rating_str, "gap_assessment")

        findings_text = (f.get("current_state") or "").strip()
        gap_delta = (f.get("gap_delta") or "").strip()
        if gap_delta:
            findings_text = f"{findings_text} {gap_delta}".strip()

        _cell_text(cells[0], f.get("clause_ref") or "",               size_pt=9, bold=True)
        _cell_text(cells[1], (f.get("control_name") or "")[:80],     size_pt=9)
        _cell_text(cells[2], (f.get("required_state") or "")[:200],  size_pt=9)
        _cell_text(cells[3], findings_text[:250],                     size_pt=9)
        _cell_text(cells[4], f.get("recommendation") or "",           size_pt=9)
        _cell_text(cells[5], colour["label"],                         size_pt=9, bold=True, colour=colour["fg"])
        _set_cell_bg(cells[5], colour["bg"])

        for ci in (0, 1, 2, 3, 4):
            _set_cell_bg(cells[ci], "FFFFFF")

    doc.add_paragraph()


def _add_gap_annex_table(doc: Document, annex_secs: dict):
    """
    Render Sterling's "ISMS (Annex A Control) Findings" as ONE continuous
    table covering all 4 Annex A groups (5/6/7/8), with a divider row
    (first cell only) between groups — matching the source template, which
    does NOT split Annex A into separate Heading-1 sections for gap assessment
    (unlike the Audit Report template).
    Columns: Control No. | Control / Clause Description | Finding | Evidence |
    Recommendation | Implementation Status.
    """
    headers = ["Control No.", "Control / Clause Description", "Finding",
               "Evidence", "Recommendation", "Implementation Status"]

    # Pre-parse findings per group so we know the total row count up front.
    group_findings: dict[int, list[dict]] = {}
    for group_num in (5, 6, 7, 8):
        sec = annex_secs.get(group_num)
        if not sec:
            continue
        content = sec.content or ""
        if content.startswith(JSON_PREFIX):
            try:
                findings = json.loads(content[len(JSON_PREFIX):])
                if isinstance(findings, list):
                    group_findings[group_num] = findings
            except json.JSONDecodeError:
                logger.warning("DOCX: Failed to parse JSON for Annex A group %s.", group_num)

    if not group_findings:
        return

    total_rows = 1 + sum(1 + len(f) for f in group_findings.values())  # header + (divider + findings) per group
    tbl = doc.add_table(rows=total_rows, cols=len(headers))
    tbl.style = "Table Grid"
    _header_row(tbl, headers, bg=DARK_NAVY)

    row_idx = 1
    for group_num in (5, 6, 7, 8):
        findings = group_findings.get(group_num)
        if findings is None:
            continue

        # Divider row — first cell only, matching Sterling's bare "5"/"6"/"7"/"8" rows
        divider_cells = tbl.rows[row_idx].cells
        _cell_text(divider_cells[0], str(group_num), size_pt=9, bold=True, colour="FFFFFF")
        _set_cell_bg(divider_cells[0], DARK_NAVY)
        for ci in range(1, len(headers)):
            _set_cell_bg(divider_cells[ci], "FFFFFF")
        row_idx += 1

        for f in findings:
            cells = tbl.rows[row_idx].cells
            rating_str = f.get("circ_rating", "Not Implemented")
            colour = get_colour_by_status(rating_str, "gap_assessment")
            description = f"{(f.get('control_name') or '').strip()} {(f.get('required_state') or '').strip()}".strip()

            _cell_text(cells[0], f.get("clause_ref") or "",            size_pt=9, bold=True)
            _cell_text(cells[1], description[:200],                     size_pt=9)
            _cell_text(cells[2], (f.get("current_state") or "")[:200], size_pt=9)
            _cell_text(cells[3], (f.get("evidence_reviewed") or "")[:150], size_pt=9)
            _cell_text(cells[4], f.get("recommendation") or "",         size_pt=9)
            _cell_text(cells[5], colour["label"],                       size_pt=9, bold=True, colour=colour["fg"])
            _set_cell_bg(cells[5], colour["bg"])

            for ci in (0, 1, 2, 3, 4):
                _set_cell_bg(cells[ci], "FFFFFF")
            row_idx += 1

    doc.add_paragraph()


# ── Gap/Non-Conformity register (from Gap DB rows) ───────────────────────────
# NOTE: not currently called by build_docx() for either report type — neither
# the Reflex Pay nor Sterling template has a section matching this layout
# (their findings tables already serve this role). Kept for potential reuse.

def _add_gap_register(doc: Document, report: Report, service_type: str):
    gaps = list(report.gaps.select_related("requirement").all())
    if not gaps:
        return

    is_gap = service_type == "gap_assessment"
    register_title = "Gap Findings Register" if is_gap else "Non-Conformity & Gap Register"
    _add_heading(doc, register_title, level=1)

    col_headers = (
        ["Ref", "Control Area", "Gap Description", "CIRC Rating", "Priority"]
        if is_gap
        else ["Ref", "Requirement", "Finding", "Audit Rating", "Severity"]
    )

    tbl = doc.add_table(rows=len(gaps) + 1, cols=len(col_headers))
    tbl.style = "Table Grid"
    _header_row(tbl, col_headers)

    for row_idx, gap in enumerate(gaps, start=1):
        rating = gap.rating or infer_rating(gap.severity, "non_compliant", service_type)
        colour = get_colour(rating, service_type)
        cells = tbl.rows[row_idx].cells

        req_code = gap.requirement.code if gap.requirement else "N/A"
        req_text = (gap.requirement.text[:70] if gap.requirement else "") or ""
        issue = gap.issue or ""
        severity = gap.severity or "medium"

        if is_gap:
            values = [
                req_code,
                req_text,
                issue,
                colour["label"],
                severity.capitalize(),
            ]
        else:
            values = [
                req_code,
                req_text,
                issue,
                colour["label"],
                severity.upper(),
            ]

        for ci, val in enumerate(values):
            _cell_text(cells[ci], val, size_pt=9,
                       bold=(ci == 3),
                       colour=(colour["fg"] if ci == 3 else None))
        _set_cell_bg(cells[3], colour["bg"])
        # Explicit white background for non-rating cells (avoid Word table-style banding)
        for ci in (0, 1, 2, 4):
            _set_cell_bg(cells[ci], "FFFFFF")

    doc.add_paragraph()


# ── Main builder ──────────────────────────────────────────────────────────────

def build_docx(report: Report) -> bytes:
    doc = Document()
    service_type = getattr(report, "service_type", "audit_report") or "audit_report"
    is_gap = service_type == "gap_assessment"
    logo = getattr(report, "logo", None)

    # Page margins
    for sect in doc.sections:
        sect.top_margin    = Cm(2.0)
        sect.bottom_margin = Cm(2.0)
        sect.left_margin   = Cm(2.5)
        sect.right_margin  = Cm(2.5)

    # ── Cover ──────────────────────────────────────────────────────────────
    _add_cover(doc, report, is_gap, logo=logo)

    # ── Document information table ─────────────────────────────────────────
    _add_doc_info_table(doc, report)
    doc.add_page_break()

    # ── Table of contents ──────────────────────────────────────────────────
    _add_toc(doc, report, is_gap)

    # ── Report sections ────────────────────────────────────────────────────
    sections = list(report.sections.order_by("order"))

    # Collect all findings across clause sections (for consolidated summary)
    all_clause_findings: list[dict] = []

    if is_gap:
        # ── Gap Assessment path — exact replica of the Sterling template ────
        # Executive Summary
        # Gap assessment Details (criteria/objectives/method/scope + info table)
        # Gap assessment Findings Definition (Status) + CIRC legend
        # Gap Assessment Detailed Findings — Management Clauses: Heading-3 per
        #     core ISO clause (4-10), each with its own 6-column table
        # ISMS (Annex A Control) Findings — ONE continuous table for all 4
        #     Annex A groups (5/6/7/8 divider rows), not separate sections
        # Summary of Findings — structured (Overall Outcome/Key Strengths/
        #     Main Gaps/Conclusion) + ratings-distribution table

        exec_summary_content = ""
        scope_content = ""
        summary_section = None
        core_secs: list = []
        annex_secs: dict = {}
        other_secs: list = []

        for sec in sections:
            name_lower = sec.section_name.lower()
            if "executive summary" in name_lower:
                exec_summary_content = sec.content or ""
                continue
            if "summary of findings" in name_lower:
                summary_section = sec
                continue
            if "conclusion" in name_lower:
                # Superseded by Summary of Findings for gap assessments — see
                # agents/pipeline.py._generate_conclusion. Skip any stale rows.
                continue
            if "scope" in name_lower and not (sec.content or "").startswith(JSON_PREFIX):
                scope_content = sec.content or ""
                continue

            kind, key = _classify_clause_section(sec.section_name)
            if kind == "core":
                core_secs.append(sec)
            elif kind == "annex":
                annex_secs[key] = sec
            else:
                other_secs.append(sec)

        # Executive Summary
        if exec_summary_content:
            _add_prose_section(doc, "Executive Summary", exec_summary_content)
            doc.add_page_break()

        # Gap assessment Details
        _add_gap_assessment_details(doc, report, scope_content)

        # Gap assessment Findings Definition (Status) + CIRC legend
        _add_gap_findings_definition(doc)

        # Gap Assessment Detailed Findings — Management Clauses
        if core_secs:
            _add_heading(doc, "Gap Assessment Detailed Findings", level=1)
            _add_heading(doc, "Management Clauses", level=2)
            for sec in core_secs:
                content = sec.content or ""
                _add_heading(doc, _format_gap_clause_heading(sec.section_name), level=3)
                if content.startswith(JSON_PREFIX):
                    raw_json = content[len(JSON_PREFIX):]
                    try:
                        findings = json.loads(raw_json)
                        if isinstance(findings, list) and findings:
                            all_clause_findings.extend(findings)
                            _add_gap_clause_table_body(doc, findings)
                            continue
                    except json.JSONDecodeError:
                        logger.warning("DOCX: Failed to parse JSON for section '%s'.", sec.section_name)
                _add_prose_section_body(doc, content)
            doc.add_page_break()

        # ISMS (Annex A Control) Findings — single combined table
        if annex_secs:
            _add_heading(doc, "ISMS (Annex A Control) Findings", level=1)
            for sec in annex_secs.values():
                content = sec.content or ""
                if content.startswith(JSON_PREFIX):
                    try:
                        findings = json.loads(content[len(JSON_PREFIX):])
                        if isinstance(findings, list):
                            all_clause_findings.extend(findings)
                    except json.JSONDecodeError:
                        pass
            _add_gap_annex_table(doc, annex_secs)
            doc.add_page_break()

        # Unclassified sections (non-ISO27001 standards) — generic fallback
        for sec in other_secs:
            content = sec.content or ""
            if content.startswith(JSON_PREFIX):
                raw_json = content[len(JSON_PREFIX):]
                try:
                    findings = json.loads(raw_json)
                    if isinstance(findings, list) and findings:
                        all_clause_findings.extend(findings)
                        _add_gap_findings_table(doc, sec.section_name, findings)
                        continue
                except json.JSONDecodeError:
                    logger.warning("DOCX: Failed to parse JSON for section '%s'.", sec.section_name)
            _add_prose_section(doc, sec.section_name, content)

        # Summary of Findings (structured) — supersedes the plain Conclusion
        # for gap assessments; Sterling has no separate top-level Conclusion
        # or Gap/Non-Conformity Register, so neither is rendered here.
        _add_gap_summary_of_findings(doc, summary_section, all_clause_findings)

    else:
        # ── Audit Report path — exact replica of the Reflex Pay template ────
        # 1.0 Executive Summary
        # 2.0 Audit Details (2.1–2.6)
        # 3.0 Audit Finding Definition — ALL core ISO clauses (4-10) merged into
        #     ONE section, each clause group as a bold sub-divider + table
        # 5./6./7./8. ISMS (Annex A Control) Audit Findings — one section per
        #     Annex A group, fixed-numbered regardless of which are present
        #     (the template itself skips "4." — intentional, not a bug)
        # 9.  Consolidated Audit Findings (clause-ref lists per rating bucket)
        # 10. Audit conclusions and audit recommendation

        # Extract named prose sections; classify the rest as core/annex/other
        exec_summary_content = ""
        scope_content = ""
        conclusion_content = ""
        core_secs: list = []
        annex_secs: dict = {}
        other_secs: list = []

        for sec in sections:
            name_lower = sec.section_name.lower()
            if "executive summary" in name_lower:
                exec_summary_content = sec.content or ""
                continue
            if "scope" in name_lower and not (sec.content or "").startswith(JSON_PREFIX):
                scope_content = sec.content or ""
                continue
            if "conclusion" in name_lower:
                conclusion_content = sec.content or ""
                continue

            kind, key = _classify_clause_section(sec.section_name)
            if kind == "core":
                core_secs.append(sec)
            elif kind == "annex":
                annex_secs[key] = sec
            else:
                other_secs.append(sec)

        # 1.0 Executive Summary
        if exec_summary_content:
            _add_prose_section(doc, "1.0  Executive Summary", exec_summary_content)
            doc.add_page_break()

        # 2.0 Audit Details §2.0–§2.6
        _add_audit_details(doc, report, scope_content)

        # 3.0 Audit Finding Definition — merged core ISO clauses (4-10)
        if core_secs:
            _add_heading(doc, "3.0  Audit Finding Definition", level=1)
            for sec in core_secs:
                content = sec.content or ""
                divider = doc.add_paragraph()
                divider.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = divider.add_run(_format_core_clause_label(sec.section_name))
                _style_run(r, size_pt=11, bold=True, colour=CORP_BLUE)

                if content.startswith(JSON_PREFIX):
                    raw_json = content[len(JSON_PREFIX):]
                    try:
                        findings = json.loads(raw_json)
                        if isinstance(findings, list) and findings:
                            all_clause_findings.extend(findings)
                            _add_audit_findings_table_body(doc, findings)
                            continue
                    except json.JSONDecodeError:
                        logger.warning("DOCX: Failed to parse JSON for section '%s'.", sec.section_name)
                _add_prose_section_body(doc, content)
            doc.add_page_break()

        # 5./6./7./8. ISMS (Annex A Control) Audit Findings — fixed numbering
        for group_num in (5, 6, 7, 8):
            sec = annex_secs.get(group_num)
            if not sec:
                continue
            content = sec.content or ""
            _add_heading(doc, "ISMS (Annex A Control) Audit Findings", level=1)
            if content.startswith(JSON_PREFIX):
                raw_json = content[len(JSON_PREFIX):]
                try:
                    findings = json.loads(raw_json)
                    if isinstance(findings, list) and findings:
                        all_clause_findings.extend(findings)
                        _add_audit_findings_table_body(doc, findings)
                except json.JSONDecodeError:
                    logger.warning("DOCX: Failed to parse JSON for section '%s'.", sec.section_name)
            else:
                _add_prose_section_body(doc, content)
            doc.add_page_break()

        # Unclassified sections (non-ISO27001 standards) — generic fallback
        section_num = 11
        for sec in other_secs:
            content = sec.content or ""
            sec_label = f"{section_num}.0  {sec.section_name}"

            if content.startswith(JSON_PREFIX):
                raw_json = content[len(JSON_PREFIX):]
                try:
                    findings = json.loads(raw_json)
                    if isinstance(findings, list) and findings:
                        all_clause_findings.extend(findings)
                        _add_audit_findings_table(doc, sec_label, findings)
                        section_num += 1
                        continue
                except json.JSONDecodeError:
                    logger.warning("DOCX: Failed to parse JSON for section '%s'.", sec.section_name)

            _add_prose_section(doc, sec_label, content)
            section_num += 1

        if other_secs:
            doc.add_page_break()

        # 9. Consolidated Audit Findings
        _add_consolidated_findings(doc, all_clause_findings)
        doc.add_page_break()

        # 10. Audit conclusions and audit recommendation
        _add_audit_conclusions(doc)

        # Append written conclusion prose if present
        if conclusion_content:
            doc.add_paragraph()
            for line in conclusion_content.split("\n"):
                line = _sanitise(line.strip())
                if not line:
                    continue
                p = doc.add_paragraph()
                r = p.add_run(line)
                _style_run(r, size_pt=10)

    # ── Apply logo to headers ──────────────────────────────────────────────
    if logo:
        if logo.placement == "every_page":
            _apply_header_logo(doc, logo)
        elif logo.placement == "selected_pages":
            _apply_header_logo_selected(doc, logo)
        # cover_only: logo is already embedded inline via _add_cover

    # ── Serialise ─────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
