"""
ISO-aware document ingestor.

Extracts text from PDFs, splits on ISO clause/control boundaries, embeds each
chunk with the ONNX embedder, and stores results in RAGDocument + RAGChunk.

Chunking strategy
-----------------
ISO 27001:2022 (requirements standard, 26 pages)
  - Split on numbered clause headers: "4", "4.1", "4.1.1", ... "10.2"
  - Also detects "Annex A" table rows (control IDs like "A.5.1")
  - Theme mapping from clause prefix:
      4 → Context, 5 → Leadership, 6 → Planning, 7 → Support,
      8 → Operation, 9 → Performance, 10 → Improvement, Annex → Annex A

ISO 27002:2022 (controls guidance standard, 164 pages)
  - Split on control headers: "5.1", "5.2", ... "8.34"
  - Theme mapping from clause prefix:
      5 → Organizational, 6 → People, 7 → Physical, 8 → Technology
  - Each control includes: statement, purpose, guidance, other information

Both:
  - Chunks shorter than MIN_CHARS are merged with the previous chunk
  - Chunks longer than MAX_CHARS are split with OVERLAP-char sliding window
  - Boilerplate (copyright, licence notices) is stripped before chunking
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import List, Tuple, Dict, Optional

logger = logging.getLogger(__name__)

# ── Chunking parameters ───────────────────────────────────────────────────────
MIN_CHARS = 150    # Merge chunks shorter than this into the previous
MAX_CHARS = 1200   # Split chunks longer than this
OVERLAP = 120      # Overlap window for long-chunk splits

# ── Regex patterns ────────────────────────────────────────────────────────────
# Top-level clause: "4 Context of the organization" or "4.1 Understanding..."
# We require the first word after the number to start with a capital letter
ISO_CLAUSE_RE = re.compile(
    r'^(\d{1,2}(?:\.\d{1,2}){0,2})\s{1,4}([A-Z][^\n]{3,80})',
    re.MULTILINE,
)

# 27002 control header — same pattern but restricted to top-level (e.g. 5.1, 8.34)
ISO_CONTROL_RE = re.compile(
    r'^([5-8]\.\d{1,2})\s{1,4}([A-Z][^\n]{3,80})',
    re.MULTILINE,
)

# Annex A control table row: "A.5.1  Policies for information security"
ANNEX_A_RE = re.compile(
    r'^(A\.\d+(?:\.\d+)?)\s{1,4}([A-Z][^\n]{3,80})',
    re.MULTILINE,
)

# Boilerplate lines to remove
BOILERPLATE_RE = re.compile(
    r'Licensed to .+?(?:\r?\n|$)'
    r'|ISO Store Order:.+?(?:\r?\n|$)'
    r'|Single user licence.+?(?:\r?\n|$)'
    r'|© ISO/IEC \d{4}.+?(?:\r?\n|$)'
    r'|Copyright protected.+?(?:\r?\n|$)',
    re.IGNORECASE,
)

# ── Theme mappings ─────────────────────────────────────────────────────────────
_27001_THEME_MAP: Dict[str, str] = {
    "0": "Introduction",
    "1": "Scope",
    "2": "Normative References",
    "3": "Terms",
    "4": "Context",
    "5": "Leadership",
    "6": "Planning",
    "7": "Support",
    "8": "Operation",
    "9": "Performance",
    "10": "Improvement",
    "A": "Annex A",
}

_27002_THEME_MAP: Dict[str, str] = {
    "5": "Organizational",
    "6": "People",
    "7": "Physical",
    "8": "Technology",
}


def _pdf_to_pages(pdf_path: str) -> Tuple[List[str], int]:
    """Return list of page texts and total page count using pdfplumber."""
    import pdfplumber  # local import — not needed at module level

    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        total = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text() or ""
            pages.append(text)
    return pages, total


def _clean(text: str) -> str:
    """Remove boilerplate and normalise whitespace."""
    text = BOILERPLATE_RE.sub("", text)
    # Collapse 3+ blank lines to 2
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _split_on_pattern(
    full_text: str,
    clause_re: re.Pattern,
) -> List[Tuple[str, str, str]]:
    """
    Split full_text into chunks at each regex match.

    Returns list of (clause_ref, section_title, body_text).
    Text before the first match becomes a preamble chunk with ref="intro".
    """
    chunks = []
    matches = list(clause_re.finditer(full_text))

    if not matches:
        return [("", "", full_text.strip())]

    # Preamble before first match
    preamble = full_text[: matches[0].start()].strip()
    if len(preamble) > MIN_CHARS:
        chunks.append(("intro", "Introduction / Front Matter", preamble))

    for i, match in enumerate(matches):
        ref = match.group(1)
        title = match.group(2).strip()
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(full_text)
        body = full_text[start:end].strip()
        chunks.append((ref, title, body))

    return chunks


def _sliding_window_split(text: str, max_chars: int, overlap: int) -> List[str]:
    """Split text into overlapping windows when it exceeds max_chars."""
    if len(text) <= max_chars:
        return [text]
    parts = []
    start = 0
    while start < len(text):
        end = start + max_chars
        parts.append(text[start:end])
        if end >= len(text):
            break
        start = end - overlap
    return parts


def _27001_theme(ref: str) -> str:
    top = ref.split(".")[0] if ref else ""
    return _27001_THEME_MAP.get(top, "General")


def _27002_theme(ref: str) -> str:
    top = ref.split(".")[0] if ref else ""
    return _27002_THEME_MAP.get(top, "General")


def _company_theme(heading: str) -> str:
    """Map a document heading to a semantic theme for company report chunks."""
    h = heading.lower()
    if any(k in h for k in ("gap", "deficien", "shortfall")):
        return "Gap Findings"
    if any(k in h for k in ("finding", "finding", "nonconform", "non-conform", "observation")):
        return "Findings"
    if any(k in h for k in ("recommend", "remediat", "roadmap", "action")):
        return "Recommendations"
    if any(k in h for k in ("risk", "threat", "vulnerab")):
        return "Risk"
    if any(k in h for k in ("scope", "boundar", "applicab")):
        return "Scope"
    if any(k in h for k in ("overview", "executive", "summary", "background")):
        return "Overview"
    if any(k in h for k in ("conclus", "close")):
        return "Conclusion"
    if any(k in h for k in ("annex", "appendix", "table")):
        return "Appendix"
    return "General"


def _infer_document_type(name: str) -> Tuple[str, str, callable]:
    """
    Return (standard_ref, doc_type, theme_fn) based on document name.
    Company report filenames (gap assessments, audit reports) are detected
    from keywords and tagged appropriately.
    """
    name_lower = name.lower()
    if "27001" in name_lower:
        # Could be a standard OR a company audit report referencing 27001
        if any(k in name_lower for k in ("audit", "assessment", "report", "gap", "finding")):
            dt = "gap_assessment" if "gap" in name_lower else "audit_report"
            return ("ISO 27001:2022", dt, _company_theme)
        return ("ISO 27001:2022", "standard", _27001_theme)
    if "27002" in name_lower:
        return ("ISO 27002:2022", "standard", _27002_theme)
    if "gap" in name_lower:
        return ("ISO 27001:2022", "gap_assessment", _company_theme)
    if "audit" in name_lower:
        return ("ISO 27001:2022", "audit_report", _company_theme)
    return ("", "standard", lambda ref: "General")


def ingest_pdf(pdf_path: str, force: bool = False) -> "RAGDocument":
    """
    Ingest a PDF into the vector store.

    1. Extract text with pdfplumber
    2. Detect ISO structure and split into semantic chunks
    3. Embed each chunk with the ONNX embedder (batch of 32)
    4. Persist RAGDocument + RAGChunk records

    If the document has already been ingested, skip unless force=True.
    Returns the RAGDocument instance.
    """
    from agents.models import RAGDocument, RAGChunk
    from agents.rag.embedder import embed

    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    name = path.stem
    standard_ref, doc_type, theme_fn = _infer_document_type(name)

    # Idempotency: skip if already ingested and not forced
    existing = RAGDocument.objects.filter(file_path=str(path)).first()
    if existing and not force:
        logger.info("Document already ingested: %s (%d chunks). Use force=True to re-ingest.",
                    name, existing.total_chunks)
        return existing
    if existing:
        logger.info("Force re-ingesting %s — deleting %d old chunks.", name, existing.total_chunks)
        existing.chunks.all().delete()
        rag_doc = existing
    else:
        rag_doc = RAGDocument.objects.create(
            name=name,
            doc_type=doc_type,
            standard_ref=standard_ref,
            file_path=str(path),
        )

    logger.info("Extracting text from %s ...", path.name)
    pages, total_pages = _pdf_to_pages(str(path))
    rag_doc.total_pages = total_pages
    rag_doc.save(update_fields=["total_pages"])

    # Join all pages with page markers so we can track page numbers
    page_texts = []
    for i, text in enumerate(pages):
        page_texts.append(f"\n\n[PAGE {i + 1}]\n{text}")
    full_text = _clean("\n".join(page_texts))

    # Choose clause pattern: 27002 controls are more specific (5.x–8.x)
    # For 27001 we also want sub-clauses (4.1.1 etc.)
    if "27002" in standard_ref:
        clause_re = ISO_CONTROL_RE
    else:
        clause_re = ISO_CLAUSE_RE

    raw_chunks = _split_on_pattern(full_text, clause_re)

    # Flatten long chunks via sliding window
    flat_chunks: List[Tuple[str, str, str]] = []
    for ref, title, body in raw_chunks:
        if len(body) > MAX_CHARS:
            sub_parts = _sliding_window_split(body, MAX_CHARS, OVERLAP)
            for j, part in enumerate(sub_parts):
                sub_ref = f"{ref}.sub{j}" if j > 0 else ref
                flat_chunks.append((sub_ref, title, part))
        elif len(body) >= MIN_CHARS:
            flat_chunks.append((ref, title, body))
        else:
            # Merge into previous chunk if too short
            if flat_chunks:
                prev_ref, prev_title, prev_body = flat_chunks[-1]
                flat_chunks[-1] = (prev_ref, prev_title, prev_body + "\n" + body)
            else:
                flat_chunks.append((ref, title, body))

    logger.info("Split into %d chunks — embedding...", len(flat_chunks))

    # Embed in batches of 32 to avoid memory spikes
    BATCH = 32
    all_chunks = []
    for batch_start in range(0, len(flat_chunks), BATCH):
        batch = flat_chunks[batch_start: batch_start + BATCH]
        texts = [body for _, _, body in batch]
        vectors = embed(texts)  # (batch, 384)

        for local_i, ((ref, title, body), vec) in enumerate(zip(batch, vectors)):
            global_i = batch_start + local_i

            # Estimate page range from [PAGE n] markers in the body
            page_nums = [int(m) for m in re.findall(r'\[PAGE (\d+)\]', body)]
            page_start = min(page_nums) if page_nums else 0
            page_end = max(page_nums) if page_nums else 0

            # Strip page markers from stored content
            clean_body = re.sub(r'\[PAGE \d+\]\n?', '', body).strip()

            all_chunks.append(RAGChunk(
                document=rag_doc,
                chunk_index=global_i,
                content=clean_body,
                clause_ref=ref,
                section_title=title,
                theme=theme_fn(ref),
                page_start=page_start,
                page_end=page_end,
                embedding=vec.tolist(),
            ))

        logger.info("  Embedded chunks %d–%d", batch_start, batch_start + len(batch) - 1)

    RAGChunk.objects.bulk_create(all_chunks, batch_size=100)

    rag_doc.total_chunks = len(all_chunks)
    rag_doc.save(update_fields=["total_chunks", "ingested_at"])

    logger.info(
        "Ingestion complete: %s — %d pages, %d chunks stored.",
        path.name, total_pages, len(all_chunks),
    )
    return rag_doc


# ── SOA Excel ingestion ────────────────────────────────────────────────────────

# Column header synonyms → canonical field names
_SOA_COL_ALIASES: Dict[str, str] = {
    # Control reference
    "control id": "control_id",
    "control ref": "control_id",
    "ref": "control_id",
    "clause": "control_id",
    "iso 27001 control": "control_id",
    "annex a control": "control_id",
    "iso 27001:2022 control": "control_id",
    # Control name / title
    "control name": "control_name",
    "control title": "control_name",
    "name": "control_name",
    "title": "control_name",
    "control": "control_name",
    "description": "control_name",
    # Applicability
    "applicable": "applicable",
    "applicability": "applicable",
    "included": "applicable",
    "include": "applicable",
    "status": "applicable",
    "in scope": "applicable",
    # Justification / reason
    "justification": "justification",
    "reason": "justification",
    "rationale": "justification",
    "justification for inclusion/exclusion": "justification",
    "inclusion/exclusion justification": "justification",
    "comments": "justification",
    "comment": "justification",
    "notes": "justification",
    "note": "justification",
    # Implementation
    "implementation": "implementation",
    "implementation status": "implementation",
    "control type": "implementation",
    "implementation notes": "implementation",
    # Responsible party
    "owner": "owner",
    "responsible": "owner",
    "control owner": "owner",
    "department": "owner",
}


def _map_soa_headers(headers: List[str]) -> Dict[int, str]:
    """Return {col_index: canonical_name} for recognised columns."""
    mapping: Dict[int, str] = {}
    for i, h in enumerate(headers):
        key = str(h).strip().lower()
        if key in _SOA_COL_ALIASES:
            canonical = _SOA_COL_ALIASES[key]
            if canonical not in mapping.values():  # first match wins
                mapping[i] = canonical
    return mapping


def _soa_row_to_text(row_data: Dict[str, str], row_num: int) -> Tuple[str, str, str]:
    """
    Convert a single SOA row dict into (clause_ref, section_title, body_text).

    Body is a human-readable narrative suitable for semantic search.
    """
    ctrl_id = row_data.get("control_id", "").strip()
    ctrl_name = row_data.get("control_name", "").strip()
    applicable = row_data.get("applicable", "").strip()
    justification = row_data.get("justification", "").strip()
    implementation = row_data.get("implementation", "").strip()
    owner = row_data.get("owner", "").strip()

    # Normalise clause ref: "A.5.1" → "5.1", "5.1" stays "5.1"
    clause_ref = re.sub(r'^A\.', '', ctrl_id) if ctrl_id else f"row{row_num}"

    # Build title
    title_raw = f"{ctrl_id} – {ctrl_name}" if ctrl_id and ctrl_name else ctrl_name or ctrl_id or f"Row {row_num}"
    title = title_raw[:250]  # section_title max_length=255

    # Truncate clause_ref to 28 chars (field max_length=30)
    clause_ref = clause_ref[:28]

    # Build narrative body
    parts = [f"Control {ctrl_id}: {ctrl_name}." if ctrl_id and ctrl_name else title + "."]
    if applicable:
        parts.append(f"Applicability: {applicable}.")
    if justification:
        parts.append(f"Justification: {justification}.")
    if implementation:
        parts.append(f"Implementation status: {implementation}.")
    if owner:
        parts.append(f"Control owner: {owner}.")

    body = " ".join(parts)
    return clause_ref, title, body


# ── DOCX company report ingestion ─────────────────────────────────────────────

def _docx_to_sections(docx_path: str) -> List[Tuple[str, str]]:
    """
    Extract (heading, body_text) sections from a .docx file.

    Splits on paragraph styles that start with "Heading" or on paragraphs that
    are short (< 120 chars), bold-only, and look like section titles.
    Also extracts tables as pipe-delimited rows appended to the preceding section.
    """
    import docx as python_docx  # local import

    doc = python_docx.Document(docx_path)
    sections: List[Tuple[str, str]] = []
    current_heading = "Front Matter"
    current_lines: List[str] = []

    def _is_heading(para) -> bool:
        style = para.style.name if para.style else ""
        if style.startswith("Heading") or style in ("Title", "Subtitle"):
            return True
        text = para.text.strip()
        if not text or len(text) > 140:
            return False
        # Heuristic: all runs bold, no terminal punctuation → looks like a heading
        runs_with_text = [r for r in para.runs if r.text.strip()]
        if runs_with_text and all(r.bold for r in runs_with_text):
            if not text.endswith((".", ",", ";")):
                return True
        # ALL CAPS short line
        if text.isupper() and len(text) > 4:
            return True
        return False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if _is_heading(para):
            if current_lines:
                sections.append((current_heading, "\n".join(current_lines)))
            current_heading = text
            current_lines = []
        else:
            current_lines.append(text)

    # Flush last section
    if current_lines:
        sections.append((current_heading, "\n".join(current_lines)))

    # Append tables as additional text blocks
    for i, table in enumerate(doc.tables):
        rows_text: List[str] = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text:
            sections.append((f"Table {i + 1}", "\n".join(rows_text)))

    return sections


def ingest_docx(docx_path: str, force: bool = False) -> "RAGDocument":
    """
    Ingest a Word (.docx) company report (gap assessment or audit report) into the
    vector store.

    Uses heading style detection to split into semantic sections.
    Each section becomes one or more RAGChunks (sliding-window if over MAX_CHARS).

    Returns the RAGDocument instance.
    """
    from agents.models import RAGDocument, RAGChunk
    from agents.rag.embedder import embed

    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    name = path.stem
    standard_ref, doc_type, theme_fn = _infer_document_type(name)

    # Idempotency
    existing = RAGDocument.objects.filter(file_path=str(path)).first()
    if existing and not force:
        logger.info("DOCX already ingested: %s (%d chunks). Use force=True to re-ingest.",
                    name, existing.total_chunks)
        return existing
    if existing:
        logger.info("Force re-ingesting DOCX %s — deleting %d old chunks.", name, existing.total_chunks)
        existing.chunks.all().delete()
        rag_doc = existing
    else:
        rag_doc = RAGDocument.objects.create(
            name=name,
            doc_type=doc_type,
            standard_ref=standard_ref or "ISO 27001:2022",
            file_path=str(path),
        )

    logger.info("Extracting sections from %s ...", path.name)
    raw_sections = _docx_to_sections(str(path))
    logger.info("  Found %d sections in %s", len(raw_sections), path.name)

    # Flatten with sliding window for long sections
    flat_chunks: List[Tuple[str, str, str]] = []
    for idx, (heading, body) in enumerate(raw_sections):
        # clause_ref = sanitised heading slug (max 28 chars for DB field)
        slug = re.sub(r'\s+', '_', heading.strip().lower())
        slug = re.sub(r'[^a-z0-9_]', '', slug)[:28] or f"sec{idx}"

        body = _clean(body)
        if len(body) > MAX_CHARS:
            sub_parts = _sliding_window_split(body, MAX_CHARS, OVERLAP)
            for j, part in enumerate(sub_parts):
                sub_ref = f"{slug}.p{j}" if j > 0 else slug
                flat_chunks.append((sub_ref[:28], heading[:250], part))
        elif len(body) >= MIN_CHARS // 2:  # company sections can be shorter
            flat_chunks.append((slug, heading[:250], body))
        else:
            if flat_chunks:
                pr, pt, pb = flat_chunks[-1]
                flat_chunks[-1] = (pr, pt, pb + "\n" + body)
            else:
                flat_chunks.append((slug, heading[:250], body))

    if not flat_chunks:
        raise ValueError(f"No text extracted from {path.name}. The document may be empty or image-only.")

    logger.info("Split into %d chunks — embedding...", len(flat_chunks))

    BATCH = 32
    all_chunks = []
    for batch_start in range(0, len(flat_chunks), BATCH):
        batch = flat_chunks[batch_start: batch_start + BATCH]
        texts = [body for _, _, body in batch]
        vectors = embed(texts)

        for local_i, ((ref, title, body), vec) in enumerate(zip(batch, vectors)):
            global_i = batch_start + local_i
            all_chunks.append(RAGChunk(
                document=rag_doc,
                chunk_index=global_i,
                content=body,
                clause_ref=ref,
                section_title=title,
                theme=theme_fn(title),
                page_start=0,
                page_end=0,
                embedding=vec.tolist(),
            ))

        logger.info("  Embedded chunks %d–%d", batch_start, batch_start + len(batch) - 1)

    RAGChunk.objects.bulk_create(all_chunks, batch_size=100)

    rag_doc.total_chunks = len(all_chunks)
    rag_doc.total_pages = len(raw_sections)  # treat sections as "pages" for display
    rag_doc.save(update_fields=["total_chunks", "total_pages", "ingested_at"])

    logger.info(
        "DOCX ingestion complete: %s (%s) — %d sections, %d chunks stored.",
        path.name, doc_type, len(raw_sections), len(all_chunks),
    )
    return rag_doc


def ingest_xlsx(xlsx_path: str, force: bool = False) -> "RAGDocument":
    """
    Ingest a Statement of Applicability (SOA) Excel workbook into the vector store.

    Each data row is converted to a semantic text chunk and embedded.
    Supports .xlsx files with flexible column naming.

    Returns the RAGDocument instance.
    """
    import openpyxl  # local import

    from agents.models import RAGDocument, RAGChunk
    from agents.rag.embedder import embed

    path = Path(xlsx_path)
    if not path.exists():
        raise FileNotFoundError(f"Excel file not found: {xlsx_path}")

    name = path.stem
    # SOA documents → doc_type="soa", standard ISO 27001:2022
    standard_ref = "ISO 27001:2022"
    doc_type = "soa"

    # Idempotency
    existing = RAGDocument.objects.filter(file_path=str(path)).first()
    if existing and not force:
        logger.info("SOA already ingested: %s (%d chunks). Use force=True to re-ingest.",
                    name, existing.total_chunks)
        return existing
    if existing:
        logger.info("Force re-ingesting SOA %s — deleting %d old chunks.", name, existing.total_chunks)
        existing.chunks.all().delete()
        rag_doc = existing
    else:
        rag_doc = RAGDocument.objects.create(
            name=name,
            doc_type=doc_type,
            standard_ref=standard_ref,
            file_path=str(path),
        )

    logger.info("Opening Excel SOA: %s ...", path.name)
    wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)

    flat_chunks: List[Tuple[str, str, str]] = []
    total_sheets = len(wb.sheetnames)

    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue

        # Find the header row — first row where we can map ≥2 columns
        header_row_idx = 0
        col_map: Dict[int, str] = {}
        for i, row in enumerate(rows[:5]):  # check first 5 rows for header
            headers = [str(c) if c is not None else "" for c in row]
            candidate = _map_soa_headers(headers)
            if len(candidate) >= 2:
                col_map = candidate
                header_row_idx = i
                break

        if not col_map:
            logger.warning("Sheet '%s': could not identify header columns — skipping.", sheet.title)
            continue

        logger.info(
            "Sheet '%s': header at row %d, mapped columns: %s",
            sheet.title, header_row_idx + 1,
            {i: v for i, v in col_map.items()},
        )

        for row_num, raw_row in enumerate(rows[header_row_idx + 1:], start=1):
            # Skip completely empty rows
            values = [c for c in raw_row if c is not None and str(c).strip()]
            if not values:
                continue

            row_data: Dict[str, str] = {}
            for col_idx, canonical in col_map.items():
                if col_idx < len(raw_row):
                    val = raw_row[col_idx]
                    row_data[canonical] = str(val).strip() if val is not None else ""

            # Skip rows without at least a control_id or control_name
            if not row_data.get("control_id") and not row_data.get("control_name"):
                continue

            clause_ref, title, body = _soa_row_to_text(row_data, row_num)
            if len(body) >= MIN_CHARS // 2:  # SOA rows are intentionally short
                flat_chunks.append((clause_ref, title, body))

    wb.close()

    if not flat_chunks:
        raise ValueError(
            f"No data rows extracted from {path.name}. "
            "Check that the sheet has recognised column headers."
        )

    logger.info("Extracted %d SOA control rows — embedding...", len(flat_chunks))

    BATCH = 32
    all_chunks = []
    for batch_start in range(0, len(flat_chunks), BATCH):
        batch = flat_chunks[batch_start: batch_start + BATCH]
        texts = [body for _, _, body in batch]
        vectors = embed(texts)

        for local_i, ((ref, title, body), vec) in enumerate(zip(batch, vectors)):
            global_i = batch_start + local_i
            all_chunks.append(RAGChunk(
                document=rag_doc,
                chunk_index=global_i,
                content=body,
                clause_ref=ref,
                section_title=title,
                theme="SOA",
                page_start=0,
                page_end=0,
                embedding=vec.tolist(),
            ))

        logger.info("  Embedded SOA chunks %d–%d", batch_start, batch_start + len(batch) - 1)

    RAGChunk.objects.bulk_create(all_chunks, batch_size=100)

    rag_doc.total_chunks = len(all_chunks)
    rag_doc.total_pages = total_sheets
    rag_doc.save(update_fields=["total_chunks", "total_pages", "ingested_at"])

    logger.info(
        "SOA ingestion complete: %s — %d sheets, %d controls stored.",
        path.name, total_sheets, len(all_chunks),
    )
    return rag_doc
